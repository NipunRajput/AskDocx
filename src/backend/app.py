import os
import uuid
import logging
from flask import Flask, request, jsonify
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_bcrypt import Bcrypt
from flask_jwt_extended import (
    JWTManager, create_access_token, jwt_required, get_jwt_identity,decode_token
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv
import requests
import PyPDF2
import docx
import json
from datetime import timedelta
# import logging # logging is already imported above
from docx.opc.exceptions import PackageNotFoundError
import docx2txt
# import docx # docx is already imported above

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart


"""AskDocx – unified backend (file processing + JWT auth)"""

# ---------------------------------------------------------------------------
# Configuration & App Setup
# ---------------------------------------------------------------------------
load_dotenv()  # Load variables from .env

app = Flask(__name__)
CORS(
    app,
    resources={r"/api/*": {"origins": "http://localhost:5173"}},
    supports_credentials=True,
    allow_headers=["Content-Type", "Authorization"],
    expose_headers=["Authorization"],
)


#Email--------------------------------------------------------------------
def send_email(to_email, subject, body):
    from_email = os.getenv("EMAIL_USER")
    password = os.getenv("EMAIL_PASS")

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        text = msg.as_string()
        server.sendmail(from_email, to_email, text)
        server.quit()
        print("Email sent successfully.")
    except Exception as e:
        print("Failed to send email:", str(e))

# Database ---------------------------------------------------------------
basedir = os.path.abspath(os.path.dirname(__file__))
app.config["SQLALCHEMY_DATABASE_URI"] = (
    "mysql+pymysql://askdocx_user:dev@127.0.0.1:3306/askdocx?charset=utf8mb4"
)


app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["JWT_SECRET_KEY"] = os.getenv("JWT_SECRET_KEY", "SET_A_REAL_SECRET")

db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
jwt = JWTManager(app)


from datetime import timedelta
app.config["JWT_ACCESS_TOKEN_EXPIRES"] = timedelta(hours=2)
# Uploads ---------------------------------------------------------------
UPLOAD_FOLDER = os.path.join(basedir, "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB

ALLOWED_EXTENSIONS = {"pdf", "docx"}

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# Groq API --------------------------------------------------------------
GROQ_API_URL = os.getenv("GROQ_API_URL")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = "meta-llama/llama-4-maverick-17b-128e-instruct" # As per your original code, ensure this model is available or update if needed

# In-memory store (replace with Redis / DB for prod)
documents_store: dict[str, dict[str, str]] = {}

# ---------------------------------------------------------------------------
# Database model:  Login Page
# ---------------------------------------------------------------------------
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    pw_hash = db.Column(db.String(128), nullable=False)

    @classmethod
    def create(cls, email: str, password: str):
        return cls(email=email, pw_hash=bcrypt.generate_password_hash(password).decode())

# Create tables on first run


# ---------------------------------------------------------------------------
# Database model  : Chat Store
# ---------------------------------------------------------------------------


class UserDocument(db.Model):
    __tablename__ = 'user_documents'  # Explicitly naming the table

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id', ondelete='CASCADE'), nullable=False)
    document_name = db.Column(db.String(255), nullable=False)
    extracted_text = db.Column(db.Text, nullable=False)  # Use db.Text; maps to TEXT/LONGTEXT
    
    # For chat history:
    # Option 1: Using SQLAlchemy's JSON type (recommended if DB supports it, e.g., MySQL 5.7.8+)
    chat = db.Column(db.JSON, nullable=True) 
    # Option 2: If DB doesn't support JSON type well, or you prefer TEXT
    # chat = db.Column(db.Text, nullable=True) # Store as a JSON string

    created_at = db.Column(db.TIMESTAMP, server_default=db.func.now())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.now(), onupdate=db.func.now())

    # Relationship to User (optional but good for ORM features)
    user = db.relationship('User', backref=db.backref('user_documents', lazy='dynamic', cascade='all, delete-orphan'))

    def __repr__(self):
        return f'<UserDocument {self.id} - {self.document_name}>'

# After defining models, ensure they are created:
# (This line is likely already in your app.py, ensure UserDocument is defined before it runs)
# with app.app_context():
#     db.create_all()


with app.app_context():
    db.create_all()

# ---------------------------------------------------------------------------
# Helper – text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(filepath: str) -> str:
    try:
        reader = PyPDF2.PdfReader(filepath)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as pdf_err:
        logging.error(f"PDF read error for {filepath}: {pdf_err}")
        raise ValueError(f"PDF read error: {pdf_err}")

def extract_text_from_docx(filepath: str) -> str:
    try:
        # 1 – try python-docx first (good for well-formed docs)
        return "\n".join(p.text for p in docx.Document(filepath).paragraphs)
    except Exception as err:
        # 2 – if **any** error occurs (metadata, corruption, etc.)
        #     fall back to docx2txt, which ignores metadata completely
        logging.warning(f"[docx] python-docx failed for {filepath} because: {err}. Falling back to docx2txt.")
        try:
            return docx2txt.process(filepath)
        except Exception as docx2txt_err:
            logging.error(f"docx2txt also failed for {filepath}: {docx2txt_err}")
            raise ValueError(f"DOCX read error with both libraries: {err}, then {docx2txt_err}")


def extract_text(filepath: str) -> str:
    ext = filepath.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(filepath)
    if ext == "docx":
        return extract_text_from_docx(filepath)
    raise ValueError(f"Unsupported file type: {ext}")

# ---------------------------------------------------------------------------
# Auth Endpoints
# ---------------------------------------------------------------------------
@app.post("/api/auth/register")
def register():
    data = request.get_json() or {}
    if not data.get("email") or not data.get("password"):
        return {"success": False, "error": "Email & password required."}, 400

    if User.query.filter_by(email=data["email"]).first():
        return {"success": False, "error": "User already exists."}, 409

    user = User.create(data["email"], data["password"])
    db.session.add(user)
    db.session.commit()
    # --- FIX: Convert user.id to string for JWT identity ---
    token = create_access_token(identity=str(user.id))
    return {"success": True, "token": token}, 201

@app.post("/api/auth/login")
def login():
    data = request.get_json() or {}
    user = User.query.filter_by(email=data.get("email")).first()
    if not user or not bcrypt.check_password_hash(user.pw_hash, data.get("password", "")): # Added default for password
        return {"success": False, "error": "Invalid credentials."}, 401
    # --- FIX: Convert user.id to string for JWT identity ---
    token = create_access_token(identity=str(user.id))
    return {"success": True, "token": token}, 200

# ---------------------------------------------------------------------------
# Document Endpoints (JWT-protected)
# ---------------------------------------------------------------------------
@app.route("/api/upload-and-process", methods=["POST"])
@jwt_required()
def upload_and_process():
    current_user_identity_str = get_jwt_identity() # This is string form of user.id
    try:
        current_user_id = int(current_user_identity_str)
    except ValueError:
        logging.error(f"Invalid user ID format in JWT: {current_user_identity_str}")
        return {"success": False, "error": "Invalid user session."}, 401

    if "document" not in request.files:
        return {"success": False, "error": "No document file part in the request"}, 400

    file = request.files["document"]
    if file.filename == "":
        return {"success": False, "error": "No selected file"}, 400

    filepath = ""
    if file and allowed_file(file.filename):
        original_filename = secure_filename(file.filename) # Ensure filename is secure
        temp_filename = f"{uuid.uuid4()}_{original_filename}"
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], temp_filename)

        try:
            file.save(filepath)
            extracted_text_content = extract_text(filepath)

            # Initial AI message for the chat history
            initial_chat_message = {
                "sender": "ai", 
                "text": f"File \"{original_filename}\" processed! Ask me anything about its content."
            }

            new_user_doc = UserDocument(
                user_id=current_user_id,
                document_name=original_filename,
                extracted_text=extracted_text_content,
                chat=[initial_chat_message] # Initialize chat history as a list with the first message
            )
            db.session.add(new_user_doc)
            db.session.commit()

            return {
                "success": True,
                "documentId": new_user_doc.id, # This is the new integer ID for the chat session
                "message": f'File "{original_filename}" processed successfully and chat session created.',
                "filename": original_filename,
            }, 200

        except ValueError as ve:
            logging.exception("Text extraction failed during upload-process")
            return {"success": False, "error": str(ve)}, 400
        except Exception as e:
            logging.exception("Upload-process failed")
            db.session.rollback() # Rollback in case of DB error during commit
            return {"success": False, "error": f"An unexpected error occurred: {str(e)}"}, 500
        finally:
            if filepath and os.path.exists(filepath):
                os.remove(filepath)
    else:
        return {"success": False, "error": "File type not allowed"}, 400

@app.route("/api/ask-question", methods=["POST"])
@jwt_required()
def ask_question():
    current_user_identity_str = get_jwt_identity()
    try:
        current_user_id = int(current_user_identity_str)
    except ValueError:
        logging.error(f"Invalid user ID format in JWT for ask-question: {current_user_identity_str}")
        return {"success": False, "error": "Invalid user session."}, 401

    data = request.get_json() or {}
    doc_session_id = data.get("documentId") # This is UserDocument.id
    question_text = data.get("question")

    if not doc_session_id or not question_text:
        return {"success": False, "error": "Missing documentId or question"}, 400

    user_doc = UserDocument.query.filter_by(id=doc_session_id, user_id=current_user_id).first()

    if not user_doc:
        return {"success": False, "error": "Document session not found or access denied"}, 404

    context_text = user_doc.extracted_text

    # --- Groq API Call (remains largely the same) ---
    prompt_messages = [
        {
            "role": "system",
            "content": (
                "You are an expert assistant. Answer strictly from the document text provided. "
                "If the answer isn't present, reply ‘The answer is not found in the provided document text.’\n\n"
                "--- Document Context Start ---\n" + context_text + "\n--- Document Context End ---"
            ),
        },
        {"role": "user", "content": question_text},
    ]

    if not GROQ_API_KEY or not GROQ_API_URL:
        logging.error("Groq API not configured.")
        return {"success": False, "error": "Groq API not configured"}, 500

    headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": GROQ_MODEL, "messages": prompt_messages, "temperature": 0.7,
        "max_tokens": 2048, "top_p": 1
    }
    logging.info(f"Sending payload to Groq for doc_session_id {doc_session_id}: {json.dumps(payload, indent=2)}")

    try:
        resp = requests.post(GROQ_API_URL, headers=headers, json=payload, timeout=60)
        if resp.status_code != 200:
            # (Include your enhanced error logging for Groq's response from previous steps here)
            error_content = resp.text
            try:
                error_json = resp.json(); detail_message = json.dumps(error_json)
            except ValueError: detail_message = error_content
            logging.error(f"Groq API Error - Status: {resp.status_code}, Detail: {detail_message}")
        resp.raise_for_status()

        ai_response_data = resp.json()
        ai_answer = ai_response_data.get("choices", [{}])[0].get("message", {}).get("content")

        if ai_answer is None:
            raise ValueError("Could not parse answer from Groq response")

        # --- Update Chat History in Database ---
        current_chat_history = list(user_doc.chat) if user_doc.chat else [] # Ensure it's a list

        current_chat_history.append({"sender": "user", "text": question_text})
        current_chat_history.append({"sender": "ai", "text": ai_answer})

        user_doc.chat = current_chat_history # Assign the modified list back
        db.session.commit()

        return {"success": True, "answer": ai_answer}, 200

    # (Keep your existing detailed exception handling for Groq API calls)
    except requests.exceptions.HTTPError as http_err:
        error_detail_from_response = "No additional error detail in response body." # ... (copy from previous)
        logging.exception(f"Groq API HTTPError. Status: {http_err.response.status_code if http_err.response else 'N/A'}. Detail: {error_detail_from_response}")
        frontend_error_message = f"Groq API request error (Status: {http_err.response.status_code if http_err.response else 'N/A'}). Details: {error_detail_from_response}"
        return {"success": False, "error": frontend_error_message}, http_err.response.status_code if http_err.response else 503
    except requests.exceptions.RequestException as req_err: # ... (copy from previous)
         logging.exception("Groq API communication (non-HTTP) failure")
         return {"success": False, "error": f"Groq API communication error: {req_err}"}, 503
    except ValueError as val_err: # ... (copy from previous)
        logging.exception("Groq API response parsing failed or other ValueError")
        db.session.rollback() # Rollback if error occurred before/during chat update
        return {"success": False, "error": str(val_err)}, 500
    except Exception as e: # ... (copy from previous)
        logging.exception("Groq API call or chat update failed with an unexpected error")
        db.session.rollback()
        return {"success": False, "error": f"An unexpected error occurred: {e}"}, 503    

# ---------------------------------------------------------------------------
# User Documents
# ---------------------------------------------------------------------------

@app.route("/api/user-documents", methods=["GET"])
@jwt_required()
def list_user_documents():
    current_user_identity_str = get_jwt_identity()
    try:
        current_user_id = int(current_user_identity_str)
    except ValueError:
        return {"success": False, "error": "Invalid user session."}, 401

    documents = UserDocument.query.filter_by(user_id=current_user_id).order_by(UserDocument.updated_at.desc()).all()

    result = [
        {
            "id": doc.id,
            "document_name": doc.document_name,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
            "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
            "preview": (doc.chat[0]['text'] if doc.chat and len(doc.chat)>0 and doc.chat[0]['sender']=='ai' 
                        else (doc.chat[1]['text'] if doc.chat and len(doc.chat)>1 and doc.chat[1]['sender']=='ai' else "Chat started."))[:100] # A short preview
        } for doc in documents
    ]
    return {"success": True, "documents": result}, 200

@app.route("/api/user-documents/<int:doc_session_id>", methods=["GET"])
@jwt_required()
def get_user_document_session(doc_session_id):
    current_user_identity_str = get_jwt_identity()
    try:
        current_user_id = int(current_user_identity_str)
    except ValueError:
        return {"success": False, "error": "Invalid user session."}, 401

    user_doc = UserDocument.query.filter_by(id=doc_session_id, user_id=current_user_id).first()

    if not user_doc:
        return {"success": False, "error": "Document session not found or access denied"}, 404

    return {
        "success": True,
        "document_session": {
            "id": user_doc.id,
            "document_name": user_doc.document_name,
            "chat_history": user_doc.chat if user_doc.chat else [], # Return chat history
            "created_at": user_doc.created_at.isoformat() if user_doc.created_at else None,
            "updated_at": user_doc.updated_at.isoformat() if user_doc.updated_at else None,
        }
    }, 200


# ---------------------------------------------------------------------------
# Forget Password
# ---------------------------------------------------------------------------

def send_email(to_email, subject, body):
    from_email = os.getenv("EMAIL_USER")
    password = os.getenv("EMAIL_PASS")

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        text = msg.as_string()
        server.sendmail(from_email, to_email, text)
        server.quit()
        print("Email sent successfully.")
    except Exception as e:
        print("Failed to send email:", str(e))

@app.route('/api/forgot-password', methods=['POST'])
def forgot_password():
    data = request.get_json()
    email = data.get('email')
    user = User.query.filter_by(email=email).first()
    if user:
        token = create_access_token(identity=user.id, expires_delta=timedelta(minutes=15))
        reset_link = f"http://localhost:5173/api/reset-password/{token}"
        send_email(user.email, "Reset your password", f"Reset link: {reset_link}")
    return jsonify({"message": "If the email is registered, a reset link will be sent."})

# ---------------------------------------------------------------------------
# Token
# ---------------------------------------------------------------------------

@app.route('/api/reset-password/<token>', methods=['POST'])
def reset_password(token):
    try:
        # Decode the token to get user ID
        decoded_token = decode_token(token)
        user_id = decoded_token['sub']  # Assuming 'sub' contains user ID

        # Get new password from request body
        data = request.get_json()
        new_password = data.get('password')

        if not new_password:
            return jsonify({"error": "Password is required"}), 400

        # Simulate updating the password in the database
        # Replace this with actual database logic
        print(f"Resetting password for user ID: {user_id} with new password: {new_password}")

        return jsonify({"message": "Password reset successful."}), 200

    except Exception as e:
        print(f"Error in reset_password: {str(e)}")
        return jsonify({"error": "Invalid or expired token."}), 400







# ---------------------------------------------------------------------------
# Misc
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    return "AskDocx backend is running!"


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s %(module)s:%(lineno)d - %(message)s"
    )
    # Ensure DB tables are created within app context
    with app.app_context():
        db.create_all()  # builds the User table if it doesn't exist

    app.run(debug=True, host="0.0.0.0", port=5000)